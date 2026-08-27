#!/usr/bin/env python3
"""Download selected papers listed in DOCX, Markdown, or XLSX files."""

from __future__ import annotations

import argparse
import csv
import html
import http.cookiejar
import ipaddress
import re
import socket
import sys
import tempfile
import urllib.error
import urllib.parse
import urllib.request
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass, field
from html.parser import HTMLParser
from pathlib import Path
from typing import Iterable


URL_RE = re.compile(r"https?://[^\s<>\]\[\)\(\"']+", re.I)
MD_LINK_RE = re.compile(r"\[([^\]]+)\]\((https?://[^\s)]+)\)", re.I)
BOLD_RE = re.compile(r"(?:\*\*|__)(.+?)(?:\*\*|__)")
RED_HTML_RE = re.compile(
    r"<(?:span|font)\b[^>]*(?:color\s*:\s*(?:red|#(?:f00|ff0000))|color\s*=\s*[\"']?(?:red|#(?:f00|ff0000)))[^>]*>(.*?)</(?:span|font)>",
    re.I | re.S,
)
TITLE_HEADERS = {"title", "paper title", "论文题目", "论文标题", "题目", "标题", "name"}
URL_HEADERS = {
    "url", "link", "paper url", "paper link", "pdf", "pdf url", "download", "download url",
    "下载地址", "论文地址", "论文链接", "全文链接", "pdf链接", "下载链接",
}
BAD_URL_HEADERS = {"code", "code url", "github", "代码", "代码地址", "代码链接", "dataset", "数据集"}
CODE_HOSTS = {"github.com", "www.github.com", "gitlab.com", "www.gitlab.com", "bitbucket.org"}
PDF_LINK_HINTS = ("pdf", "download paper", "download pdf", "full text", "view paper", "view pdf")
PDF_LINK_EXCLUDES = ("supp", "appendix", "poster", "slide", "presentation", "dataset", "code")
MAX_LANDING_HOPS = 3


@dataclass
class PaperItem:
    title: str
    url: str
    source_location: str
    red: bool = False
    bold: bool = False
    selected: bool = False
    status: str = "not selected"
    saved_file: str = ""
    final_url: str = ""
    reason: str = ""


def clean_text(value: object) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def normalize_header(value: object) -> str:
    return clean_text(value).lower().rstrip(":：")


def urls_in(text: str) -> list[str]:
    return [u.rstrip(".,;:，。；：") for u in URL_RE.findall(text or "")]


def is_red_rgb(rgb: str | None) -> bool:
    if not rgb:
        return False
    value = str(rgb).upper().replace("#", "")[-6:]
    if len(value) != 6 or not re.fullmatch(r"[0-9A-F]{6}", value):
        return False
    r, g, b = int(value[:2], 16), int(value[2:4], 16), int(value[4:], 16)
    return r >= 150 and r >= g + 55 and r >= b + 55


def looks_like_paper_url(url: str) -> bool:
    host = (urllib.parse.urlparse(url).hostname or "").lower()
    return host not in CODE_HOSTS


def choose_url(candidates: Iterable[str]) -> str:
    unique = list(dict.fromkeys(u for u in candidates if u))
    papers = [u for u in unique if looks_like_paper_url(u)]
    direct = [u for u in papers if ".pdf" in urllib.parse.urlparse(u).path.lower()]
    return (direct or papers or [""])[0]


def deduplicate(items: list[PaperItem]) -> list[PaperItem]:
    seen: set[tuple[str, str]] = set()
    output: list[PaperItem] = []
    for item in items:
        key = (item.title.casefold(), item.url.casefold())
        if key not in seen and (item.title or item.url):
            seen.add(key)
            output.append(item)
    return output


def docx_run_red(run) -> bool:
    color = getattr(getattr(run, "font", None), "color", None)
    rgb = getattr(color, "rgb", None)
    return is_red_rgb(str(rgb) if rgb else None)


def docx_run_bold(run) -> bool:
    if run.bold is True:
        return True
    style = getattr(run, "style", None)
    return bool(style and getattr(getattr(style, "font", None), "bold", False))


def docx_para_meta(paragraph) -> tuple[bool, bool]:
    runs = [r for r in paragraph.runs if clean_text(r.text)]
    return any(docx_run_red(r) for r in runs), any(docx_run_bold(r) for r in runs)


def docx_hyperlinks(element, part) -> list[str]:
    found: list[str] = []
    rel_ns = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"
    for link in element._element.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink"):
        rid = link.get(rel_ns)
        if rid and rid in part.rels:
            found.append(part.rels[rid].target_ref)
    return found


def parse_docx(path: Path) -> list[PaperItem]:
    try:
        from docx import Document
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except ImportError as exc:
        raise RuntimeError("DOCX support requires python-docx") from exc

    doc = Document(path)
    items: list[PaperItem] = []

    for table_index, table in enumerate(doc.tables, 1):
        if not table.rows:
            continue
        headers = [normalize_header(c.text) for c in table.rows[0].cells]
        title_col = next((i for i, h in enumerate(headers) if h in TITLE_HEADERS), None)
        url_col = next((i for i, h in enumerate(headers) if h in URL_HEADERS and h not in BAD_URL_HEADERS), None)
        if title_col is None:
            continue
        for row_index, row in enumerate(table.rows[1:], 2):
            if title_col >= len(row.cells):
                continue
            title_cell = row.cells[title_col]
            title = clean_text(title_cell.text)
            red = bold = False
            for paragraph in title_cell.paragraphs:
                p_red, p_bold = docx_para_meta(paragraph)
                red, bold = red or p_red, bold or p_bold
            candidates: list[str] = []
            cells = [row.cells[url_col]] if url_col is not None and url_col < len(row.cells) else row.cells
            for cell in cells:
                candidates.extend(urls_in(cell.text))
                for paragraph in cell.paragraphs:
                    candidates.extend(docx_hyperlinks(paragraph, paragraph.part))
            items.append(PaperItem(title, choose_url(candidates), f"table {table_index}, row {row_index}", red, bold))

    paragraphs = list(doc.paragraphs)
    for index, paragraph in enumerate(paragraphs):
        text = clean_text(paragraph.text)
        if not text:
            continue
        candidates = urls_in(text) + docx_hyperlinks(paragraph, paragraph.part)
        red, bold = docx_para_meta(paragraph)
        title = URL_RE.sub("", text).strip(" -–—:：")
        if candidates and title:
            items.append(PaperItem(title, choose_url(candidates), f"paragraph {index + 1}", red, bold))
        elif (red or bold) and index + 1 < len(paragraphs):
            next_p = paragraphs[index + 1]
            next_urls = urls_in(next_p.text) + docx_hyperlinks(next_p, next_p.part)
            if next_urls:
                items.append(PaperItem(text, choose_url(next_urls), f"paragraphs {index + 1}-{index + 2}", red, bold))
    return deduplicate(items)


def excel_color_red(font) -> bool:
    color = getattr(font, "color", None)
    if not color:
        return False
    if color.type == "rgb":
        return is_red_rgb(color.rgb)
    return color.type == "indexed" and color.indexed in {2, 10}


def excel_cell_urls(cell) -> list[str]:
    output = urls_in(clean_text(cell.value))
    if cell.hyperlink and cell.hyperlink.target:
        output.append(cell.hyperlink.target)
    return output


def parse_xlsx(path: Path, sheet_name: str | None) -> list[PaperItem]:
    try:
        import openpyxl
    except ImportError as exc:
        raise RuntimeError("XLSX support requires openpyxl") from exc

    workbook = openpyxl.load_workbook(path, data_only=True, read_only=False)
    if sheet_name:
        if sheet_name not in workbook.sheetnames:
            raise RuntimeError(f"Excel sheet not found: {sheet_name}")
        sheets = [workbook[sheet_name]]
    else:
        sheets = [s for s in workbook.worksheets if s.sheet_state == "visible"]
    items: list[PaperItem] = []
    for sheet in sheets:
        rows = list(sheet.iter_rows())
        if not rows:
            continue
        headers = [normalize_header(c.value) for c in rows[0]]
        title_col = next((i for i, h in enumerate(headers) if h in TITLE_HEADERS), None)
        url_col = next((i for i, h in enumerate(headers) if h in URL_HEADERS and h not in BAD_URL_HEADERS), None)
        if title_col is None:
            continue
        for row_index, row in enumerate(rows[1:], 2):
            if title_col >= len(row):
                continue
            cell = row[title_col]
            title = clean_text(cell.value)
            if not title:
                continue
            candidates: list[str] = []
            cells = [row[url_col]] if url_col is not None and url_col < len(row) else row
            for candidate_cell in cells:
                candidates.extend(excel_cell_urls(candidate_cell))
            items.append(PaperItem(
                title, choose_url(candidates), f"sheet {sheet.title}, row {row_index}",
                excel_color_red(cell.font), bool(cell.font.bold),
            ))
    workbook.close()
    return deduplicate(items)


def strip_markdown(text: str) -> str:
    text = re.sub(r"<[^>]+>", "", text)
    text = re.sub(r"(?:\*\*|__)(.+?)(?:\*\*|__)", r"\1", text)
    text = re.sub(r"[`*_~]", "", text)
    return clean_text(html.unescape(text))


def markdown_meta(text: str) -> tuple[bool, bool]:
    return bool(RED_HTML_RE.search(text)), bool(BOLD_RE.search(text))


def parse_markdown(path: Path) -> list[PaperItem]:
    lines = path.read_text(encoding="utf-8-sig").splitlines()
    items: list[PaperItem] = []
    index = 0
    while index < len(lines):
        raw = lines[index].strip()
        if not raw:
            index += 1
            continue
        # Markdown tables: parse header, separator, and data rows as a unit.
        if "|" in raw and index + 1 < len(lines) and re.match(r"^\s*\|?\s*:?-{3,}", lines[index + 1]):
            headers = [normalize_header(c) for c in raw.strip("|").split("|")]
            title_col = next((i for i, h in enumerate(headers) if h in TITLE_HEADERS), None)
            url_col = next((i for i, h in enumerate(headers) if h in URL_HEADERS and h not in BAD_URL_HEADERS), None)
            index += 2
            while index < len(lines) and "|" in lines[index]:
                cells = [c.strip() for c in lines[index].strip().strip("|").split("|")]
                if title_col is not None and title_col < len(cells):
                    title_raw = cells[title_col]
                    red, bold = markdown_meta(title_raw)
                    candidates: list[str] = []
                    source_cells = [cells[url_col]] if url_col is not None and url_col < len(cells) else cells
                    for value in source_cells:
                        candidates.extend(u for _, u in MD_LINK_RE.findall(value))
                        candidates.extend(urls_in(value))
                    items.append(PaperItem(strip_markdown(title_raw), choose_url(candidates), f"line {index + 1}", red, bold))
                index += 1
            continue

        links = MD_LINK_RE.findall(raw)
        candidates = [u for _, u in links] + urls_in(raw)
        red, bold = markdown_meta(raw)
        title_raw = MD_LINK_RE.sub(lambda m: m.group(1), raw)
        title_raw = URL_RE.sub("", title_raw).strip(" -–—:：#")
        if candidates and title_raw:
            items.append(PaperItem(strip_markdown(title_raw), choose_url(candidates), f"line {index + 1}", red, bold))
        elif (red or bold) and index + 1 < len(lines):
            next_raw = lines[index + 1]
            next_candidates = [u for _, u in MD_LINK_RE.findall(next_raw)] + urls_in(next_raw)
            if next_candidates:
                items.append(PaperItem(strip_markdown(raw.strip("# -")), choose_url(next_candidates), f"lines {index + 1}-{index + 2}", red, bold))
                index += 1
        index += 1
    return deduplicate(items)


def parse_input(path: Path, sheet_name: str | None) -> list[PaperItem]:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return parse_docx(path)
    if suffix == ".md":
        return parse_markdown(path)
    if suffix == ".xlsx":
        return parse_xlsx(path, sheet_name)
    raise RuntimeError("Unsupported input format. Use .docx, .md, or .xlsx")


def apply_selection(items: list[PaperItem], mode: str, title_regex: str | None) -> None:
    pattern = re.compile(title_regex, re.I) if title_regex else None
    for item in items:
        formatted = {
            "all": True,
            "red": item.red,
            "bold": item.bold,
            "marked": item.red or item.bold,
            "red-and-bold": item.red and item.bold,
        }[mode]
        item.selected = formatted and (not pattern or bool(pattern.search(item.title)))
        item.status = "planned" if item.selected else "not selected"
        if item.selected and not item.url:
            item.status = "failed"
            item.reason = "No paper/download URL found in the source row or entry"


def validate_public_url(url: str) -> None:
    parsed = urllib.parse.urlparse(url)
    if parsed.scheme not in {"http", "https"} or not parsed.hostname:
        raise ValueError("Only HTTP(S) URLs with a host are allowed")
    host = parsed.hostname.lower()
    if host == "localhost" or host.endswith(".local"):
        raise ValueError("Local network URLs are not allowed")
    try:
        infos = socket.getaddrinfo(host, parsed.port or (443 if parsed.scheme == "https" else 80), type=socket.SOCK_STREAM)
    except socket.gaierror as exc:
        raise ValueError(f"DNS lookup failed: {exc}") from exc
    for info in infos:
        address = ipaddress.ip_address(info[4][0])
        if address.is_private or address.is_loopback or address.is_link_local or address.is_reserved or address.is_multicast or address.is_unspecified:
            raise ValueError("Local, private, or reserved network targets are not allowed")


class SafeRedirectHandler(urllib.request.HTTPRedirectHandler):
    def redirect_request(self, req, fp, code, msg, headers, newurl):
        validate_public_url(newurl)
        return super().redirect_request(req, fp, code, msg, headers, newurl)


def arxiv_pdf_url(url: str) -> str:
    parsed = urllib.parse.urlparse(url)
    if parsed.hostname in {"arxiv.org", "www.arxiv.org"} and parsed.path.startswith("/abs/"):
        paper_id = parsed.path.removeprefix("/abs/")
        return urllib.parse.urlunparse(("https", "arxiv.org", f"/pdf/{paper_id}.pdf", "", "", ""))
    return url


def cvf_pdf_url(url: str) -> str:
    """Resolve the stable CVF paper-page pattern without depending on button text."""
    parsed = urllib.parse.urlparse(url)
    if (parsed.hostname or "").lower() not in {"openaccess.thecvf.com", "www.openaccess.thecvf.com"}:
        return ""
    if "/html/" not in parsed.path or not parsed.path.endswith("_paper.html"):
        return ""
    path = parsed.path.replace("/html/", "/papers/", 1)[:-5] + ".pdf"
    return urllib.parse.urlunparse((parsed.scheme or "https", parsed.netloc, path, "", "", ""))


class PdfCandidateParser(HTMLParser):
    """Collect PDF-like links from anchors, metadata, embedded viewers, and buttons."""

    def __init__(self, base_url: str):
        super().__init__(convert_charrefs=True)
        self.base_url = base_url
        self.candidates: list[tuple[int, int, str]] = []
        self._order = 0
        self._anchor_href = ""
        self._anchor_text: list[str] = []

    def add(self, raw_url: str, score: int, context: str = "") -> None:
        raw_url = html.unescape(clean_text(raw_url))
        if not raw_url or raw_url.startswith(("javascript:", "mailto:", "#")):
            return
        absolute = urllib.parse.urljoin(self.base_url, raw_url)
        combined = f"{raw_url} {context}".lower()
        if any(bad in combined for bad in PDF_LINK_EXCLUDES):
            return
        path = urllib.parse.urlparse(absolute).path.lower()
        if ".pdf" in path:
            score += 8
        if score <= 0:
            return
        self.candidates.append((score, -self._order, absolute))
        self._order += 1

    def handle_starttag(self, tag: str, attrs) -> None:
        attributes = {str(k).lower(): str(v or "") for k, v in attrs}
        tag = tag.lower()
        if tag == "a":
            self._anchor_href = attributes.get("href", "")
            self._anchor_text = []
        elif tag == "meta":
            key = (attributes.get("name") or attributes.get("property")).lower()
            if key in {"citation_pdf_url", "citation_pdf", "dc.source", "og:pdf"}:
                self.add(attributes.get("content", ""), 20, key)
        elif tag == "link":
            relation = f"{attributes.get('rel', '')} {attributes.get('type', '')}".lower()
            if "pdf" in relation or "alternate" in relation:
                self.add(attributes.get("href", ""), 12, relation)
        elif tag in {"iframe", "embed", "object"}:
            self.add(attributes.get("src") or attributes.get("data", ""), 10, tag)
        elif tag in {"button", "input"}:
            context = " ".join([attributes.get("aria-label", ""), attributes.get("title", ""), attributes.get("value", "")])
            for key in ("data-href", "data-url", "formaction", "href"):
                if attributes.get(key):
                    self.add(attributes[key], 8 if "pdf" in context.lower() else 2, context)
            onclick = attributes.get("onclick", "")
            match = re.search(r"(?:location(?:\.href)?|window\.open)\s*\(?\s*[=:]?\s*[\"']([^\"']+)", onclick, re.I)
            if match:
                self.add(match.group(1), 8 if "pdf" in context.lower() else 2, context)

    def handle_data(self, data: str) -> None:
        if self._anchor_href:
            self._anchor_text.append(data)

    def handle_endtag(self, tag: str) -> None:
        if tag.lower() != "a" or not self._anchor_href:
            return
        label = clean_text(" ".join(self._anchor_text)).lower().strip("[]()")
        score = 0
        if label == "pdf":
            score += 20
        elif any(hint in label for hint in PDF_LINK_HINTS):
            score += 8
        self.add(self._anchor_href, score, label)
        self._anchor_href = ""
        self._anchor_text = []


def extract_pdf_link(page: bytes, base_url: str) -> str:
    known_cvf = cvf_pdf_url(base_url)
    if known_cvf:
        return known_cvf
    text = page.decode("utf-8", errors="ignore")
    parser = PdfCandidateParser(base_url)
    try:
        parser.feed(text)
        parser.close()
    except (ValueError, AssertionError):
        pass
    return max(parser.candidates, default=(0, 0, ""))[2]


def open_url(opener, url: str, timeout: float, referer: str = ""):
    validate_public_url(url)
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 Chrome/124 Safari/537.36",
        "Accept": "application/pdf,text/html;q=0.9,*/*;q=0.5",
        "Accept-Language": "en-US,en;q=0.8",
    }
    if referer:
        headers["Referer"] = referer
    request = urllib.request.Request(url, headers=headers)
    return opener.open(request, timeout=timeout)


def safe_filename(title: str) -> str:
    value = re.sub(r"[<>:\"/\\|?*\x00-\x1F]", "_", title).strip(" .")
    value = re.sub(r"\s+", " ", value)[:170]
    return (value or "paper") + ".pdf"


def available_path(output_dir: Path, filename: str, overwrite: bool) -> tuple[Path, bool]:
    target = output_dir / filename
    if overwrite or not target.exists():
        return target, False
    return target, True


def resolve_and_download(opener, start_url: str, target: Path, timeout: float, max_bytes: int) -> tuple[Path, str, int]:
    """Follow a bounded chain of landing pages until a verified PDF is reached."""
    current_url = start_url
    referer = ""
    visited: set[str] = set()
    for hop in range(MAX_LANDING_HOPS + 1):
        if current_url in visited:
            raise ValueError("Landing-page link cycle detected")
        visited.add(current_url)
        with open_url(opener, current_url, timeout, referer) as response:
            final_url = response.geturl()
            content_type = response.headers.get_content_type().lower()
            content_length = response.headers.get("Content-Length")
            if content_length and int(content_length) > max_bytes:
                raise ValueError("File exceeds configured maximum size")
            prefix = response.read(1024)
            if b"%PDF-" in prefix:
                return write_stream(target, prefix, response, max_bytes), final_url, hop
            if content_type not in {"text/html", "application/xhtml+xml"} and b"<html" not in prefix.lower() and b"<!doctype" not in prefix.lower():
                raise ValueError(f"Response is not a PDF or HTML landing page (content type: {content_type})")
            page_limit = min(6 * 1024 * 1024, max_bytes)
            page = prefix + response.read(page_limit - len(prefix))
            pdf_url = extract_pdf_link(page, final_url)
            if not pdf_url:
                raise ValueError("Landing page did not expose a PDF link, PDF button, or citation_pdf_url metadata")
            referer, current_url = final_url, pdf_url
    raise ValueError(f"PDF was not reached after {MAX_LANDING_HOPS} landing-page hops")


def download_one(item: PaperItem, output_dir: Path, timeout: float, max_bytes: int, overwrite: bool) -> PaperItem:
    if not item.selected or item.status == "failed":
        return item
    target, exists = available_path(output_dir, safe_filename(item.title), overwrite)
    if exists:
        item.status = "skipped"
        item.saved_file = str(target)
        item.reason = "File already exists; overwrite was not enabled"
        return item
    cookie_jar = http.cookiejar.CookieJar()
    opener = urllib.request.build_opener(SafeRedirectHandler(), urllib.request.HTTPCookieProcessor(cookie_jar))
    url = arxiv_pdf_url(item.url)
    temp_path: Path | None = None
    try:
        temp_path, final_url, landing_hops = resolve_and_download(opener, url, target, timeout, max_bytes)
        if not temp_path or temp_path.stat().st_size < 1024:
            raise ValueError("Downloaded PDF is empty or truncated")
        temp_path.replace(target)
        item.status = "success"
        item.saved_file = str(target)
        item.final_url = final_url
        if landing_hops:
            item.reason = f"Resolved PDF through {landing_hops} landing page(s)"
    except (urllib.error.HTTPError, urllib.error.URLError, TimeoutError, ValueError, OSError) as exc:
        if temp_path and temp_path.exists():
            temp_path.unlink(missing_ok=True)
        item.status = "failed"
        item.reason = clean_text(exc)
    return item


def write_stream(target: Path, prefix: bytes, response, max_bytes: int) -> Path:
    handle = tempfile.NamedTemporaryFile(prefix=target.stem + "-", suffix=".part", dir=target.parent, delete=False)
    path = Path(handle.name)
    total = 0
    try:
        with handle:
            handle.write(prefix)
            total += len(prefix)
            while True:
                chunk = response.read(64 * 1024)
                if not chunk:
                    break
                total += len(chunk)
                if total > max_bytes:
                    raise ValueError("Download exceeded configured maximum size")
                handle.write(chunk)
    except Exception:
        path.unlink(missing_ok=True)
        raise
    return path


def report_counts(items: list[PaperItem]) -> dict[str, int]:
    return {
        "total": len(items),
        "selected": sum(i.selected for i in items),
        "success": sum(i.status == "success" for i in items),
        "failed": sum(i.status == "failed" for i in items),
        "skipped": sum(i.status == "skipped" for i in items),
        "planned": sum(i.status == "planned" for i in items),
        "not_selected": sum(i.status == "not selected" for i in items),
    }


def write_reports(items: list[PaperItem], output_dir: Path, prefix: str, source: Path, mode: str, dry_run: bool) -> tuple[Path, Path]:
    csv_path = output_dir / f"{prefix}.csv"
    md_path = output_dir / f"{prefix}.md"
    fields = ["title", "url", "source_location", "red", "bold", "selected", "status", "saved_file", "final_url", "reason"]
    with csv_path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fields)
        writer.writeheader()
        for item in items:
            writer.writerow({field: getattr(item, field) for field in fields})
    counts = report_counts(items)
    lines = [
        "# Paper download report", "", f"- Source: `{source}`", f"- Selection: `{mode}`",
        f"- Dry run: `{str(dry_run).lower()}`", f"- Total entries: {counts['total']}",
        f"- Selected: {counts['selected']}", f"- Successful: {counts['success']}",
        f"- Failed: {counts['failed']}", f"- Skipped: {counts['skipped']}",
        f"- Planned: {counts['planned']}", f"- Not selected: {counts['not_selected']}", "",
        "| Status | Title | Source location | Detail |", "|---|---|---|---|",
    ]
    for item in items:
        if item.status == "success" and item.reason:
            detail = f"{item.saved_file} — {item.reason}; final URL: {item.final_url}"
        else:
            detail = item.reason or item.saved_file or item.url
        escaped_title = item.title.replace("|", "\\|")
        escaped_source = item.source_location.replace("|", "\\|")
        escaped_detail = clean_text(detail).replace("|", "\\|")
        lines.append(f"| {item.status} | {escaped_title} | {escaped_source} | {escaped_detail} |")
    md_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return md_path, csv_path


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input", type=Path, help="Input .docx, .md, or .xlsx paper list")
    parser.add_argument("--output-dir", type=Path, required=True, help="Directory for PDFs and reports")
    parser.add_argument("--select", choices=["all", "red", "bold", "marked", "red-and-bold"], default="all")
    parser.add_argument("--sheet", help="One Excel sheet name; default is all visible sheets")
    parser.add_argument("--title-regex", help="Additional case-insensitive title regex")
    parser.add_argument("--workers", type=int, default=4)
    parser.add_argument("--timeout", type=float, default=30.0)
    parser.add_argument("--max-mb", type=int, default=200)
    parser.add_argument("--overwrite", action="store_true")
    parser.add_argument("--dry-run", action="store_true")
    parser.add_argument("--report-prefix", default="paper_download_report")
    return parser


def main(argv: list[str] | None = None) -> int:
    args = build_parser().parse_args(argv)
    source = args.input.resolve()
    if not source.is_file():
        print(f"ERROR: input file not found: {source}", file=sys.stderr)
        return 2
    if not 1 <= args.workers <= 16 or args.timeout <= 0 or args.max_mb <= 0:
        print("ERROR: workers must be 1-16, timeout and max-mb must be positive", file=sys.stderr)
        return 2
    try:
        items = parse_input(source, args.sheet)
        apply_selection(items, args.select, args.title_regex)
    except (RuntimeError, re.error, OSError) as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 2
    output_dir = args.output_dir.resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    if not args.dry_run:
        selected = [i for i in items if i.selected and i.status != "failed"]
        with ThreadPoolExecutor(max_workers=args.workers) as pool:
            futures = {
                pool.submit(download_one, item, output_dir, args.timeout, args.max_mb * 1024 * 1024, args.overwrite): item
                for item in selected
            }
            for future in as_completed(futures):
                future.result()
    md_path, csv_path = write_reports(items, output_dir, args.report_prefix, source, args.select, args.dry_run)
    counts = report_counts(items)
    print(f"Parsed={counts['total']} Selected={counts['selected']} Success={counts['success']} Failed={counts['failed']} Skipped={counts['skipped']} Planned={counts['planned']}")
    print(f"Markdown report: {md_path}")
    print(f"CSV report: {csv_path}")
    return 1 if counts["failed"] else 0


if __name__ == "__main__":
    raise SystemExit(main())
